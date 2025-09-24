"""
📄 Document Processor Agent - Procesa y extrae información de documentos
"""
import os
import re
import pandas as pd
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from datetime import datetime

# Imports para procesamiento de documentos
import PyPDF2
import docx
from PIL import Image
import pytesseract  # OCR

# Import del agente base
from .base_agent import BaseAgent, AgentTask, TaskPriority, validate_agent_config
from ..core.config import AgentConfig
from ..services.ai_service import AIService

@validate_agent_config(["supported_formats", "max_file_size"])
class DocumentProcessorAgent(BaseAgent):
    """
    🤖 Agente especializado en procesamiento de documentos
    
    Capacidades:
    - Extraer texto de PDFs, Word, imágenes
    - OCR de imágenes y documentos escaneados
    - Análisis de contenido con IA
    - Extracción de datos estructurados
    - Clasificación automática de documentos
    """
    
    def __init__(self, name: str, description: str, config: Optional[Dict[str, Any]] = None):
        # Configuración por defecto
        default_config = {
            **AgentConfig.DOCUMENT_PROCESSOR,
            "ocr_enabled": True,
            "ai_analysis_enabled": True,
            "auto_classify": True,
            "extract_entities": True
        }
        
        if config:
            default_config.update(config)
        
        super().__init__(name, description, default_config)
        
        # Inicializar servicio de IA
        self.ai_service = AIService()
        
        # Estadísticas específicas del procesador
        self.processed_documents = 0
        self.extracted_pages = 0
        self.ocr_operations = 0
        
        self.logger.info("📄 Document Processor Agent inicializado")
    
    def get_capabilities(self) -> List[str]:
        """Capacidades del agente procesador de documentos"""
        return [
            "extract_text_pdf",
            "extract_text_docx", 
            "extract_text_image_ocr",
            "classify_document",
            "extract_entities",
            "analyze_sentiment",
            "extract_structured_data",
            "generate_summary",
            "detect_language"
        ]
    
    async def process_task(self, task: AgentTask) -> Any:
        """Procesar tarea específica del documento"""
        task_name = task.name
        task_data = task.data
        
        self.logger.info(f"🔄 Procesando: {task_name}")
        
        try:
            if task_name == "extract_text":
                return await self._extract_text(task_data)
            
            elif task_name == "analyze_document":
                return await self._analyze_document(task_data)
            
            elif task_name == "classify_document":
                return await self._classify_document(task_data)
            
            elif task_name == "extract_entities":
                return await self._extract_entities(task_data)
            
            elif task_name == "process_batch":
                return await self._process_batch(task_data)
            
            elif task_name == "extract_tables":
                return await self._extract_tables(task_data)
            
            else:
                raise ValueError(f"Tarea no soportada: {task_name}")
                
        except Exception as e:
            self.logger.error(f"❌ Error procesando {task_name}: {str(e)}")
            raise
    
    # 📄 Métodos de extracción de texto
    
    async def _extract_text(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Extraer texto de un documento"""
        file_path = data.get("file_path")
        
        if not file_path or not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        result = {
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "file_size": os.path.getsize(file_path),
            "processed_at": datetime.now().isoformat(),
            "text": "",
            "metadata": {},
            "pages": 0,
            "word_count": 0
        }
        
        try:
            if file_extension == ".pdf":
                result.update(await self._extract_from_pdf(file_path))
            
            elif file_extension == ".docx":
                result.update(await self._extract_from_docx(file_path))
            
            elif file_extension in [".txt"]:
                result.update(await self._extract_from_txt(file_path))
            
            elif file_extension in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
                result.update(await self._extract_from_image(file_path))
            
            else:
                raise ValueError(f"Formato no soportado: {file_extension}")
            
            # Calcular estadísticas
            result["word_count"] = len(result["text"].split())
            result["char_count"] = len(result["text"])
            
            # Detectar idioma si la IA está habilitada
            if self.config.get("ai_analysis_enabled"):
                result["language"] = await self._detect_language(result["text"])
            
            self.processed_documents += 1
            self.logger.success(f"✅ Texto extraído: {result['word_count']} palabras")
            
            return result
            
        except Exception as e:
            self.logger.error(f"❌ Error extrayendo texto: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extraer texto de PDF"""
        text = ""
        pages = 0
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages = len(pdf_reader.pages)
                
                # Extraer metadata
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', '')
                    }
                
                # Extraer texto de todas las páginas
                for page_num in range(pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            self.extracted_pages += pages
            
        except Exception as e:
            # Si falla la extracción normal, intentar OCR
            if self.config.get("ocr_enabled"):
                self.logger.warning(f"⚠️ Extracción PDF falló, intentando OCR: {str(e)}")
                return await self._extract_with_ocr(file_path)
            else:
                raise e
        
        return {
            "text": text.strip(),
            "pages": pages,
            "metadata": metadata,
            "extraction_method": "pypdf2"
        }
    
    async def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extraer texto de documento Word"""
        try:
            doc = docx.Document(file_path)
            
            # Extraer texto de párrafos
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            text = "\n".join(paragraphs)
            
            # Extraer texto de tablas
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            if tables_text:
                text += "\n\nTABLAS:\n" + "\n".join(tables_text)
            
            # Metadata básica
            metadata = {
                "paragraphs_count": len(paragraphs),
                "tables_count": len(doc.tables)
            }
            
            return {
                "text": text.strip(),
                "pages": 1,  # Word no tiene concepto de páginas como PDF
                "metadata": metadata,
                "extraction_method": "docx"
            }
            
        except Exception as e:
            self.logger.error(f"❌ Error extrayendo de DOCX: {str(e)}")
            raise
    
    async def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extraer texto de archivo de texto plano"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text.strip(),
                "pages": 1,
                "metadata": {"encoding": "utf-8"},
                "extraction_method": "plain_text"
            }
            
        except UnicodeDecodeError:
            # Intentar con otras codificaciones
            encodings = ['latin-1', 'cp1252', 'ascii']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    return {
                        "text": text.strip(),
                        "pages": 1,
                        "metadata": {"encoding": encoding},
                        "extraction_method": "plain_text"
                    }
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("No se pudo decodificar el archivo de texto")
    
    async def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extraer texto de imagen usando OCR"""
        if not self.config.get("ocr_enabled"):
            raise ValueError("OCR no está habilitado")
        
        try:
            # Abrir imagen
            image = Image.open(file_path)
            
            # Extraer texto con OCR
            text = pytesseract.image_to_string(image, lang='spa+eng')
            
            # Obtener información adicional
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidence = sum([int(conf) for conf in data['conf'] if int(conf) > 0]) / len([conf for conf in data['conf'] if int(conf) > 0])
            
            metadata = {
                "image_size": image.size,
                "image_mode": image.mode,
                "ocr_confidence": round(confidence, 2) if confidence > 0 else 0
            }
            
            self.ocr_operations += 1
            
            return {
                "text": text.strip(),
                "pages": 1,
                "metadata": metadata,
                "extraction_method": "ocr"
            }
            
        except Exception as e:
            self.logger.error(f"❌ Error en OCR: {str(e)}")
            raise
    
    async def _extract_with_ocr(self, file_path: str) -> Dict[str, Any]:
        """Extraer usando OCR como fallback para PDFs problemáticos"""
        # Nota: Esto requeriría convertir PDF a imágenes primero
        # Por simplicidad, retornamos error por ahora
        raise NotImplementedError("OCR fallback para PDF no implementado aún")
    
    # 🧠 Métodos de análisis con IA
    
    async def _analyze_document(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Análisis completo del documento con IA"""
        text = data.get("text", "")
        
        if not text.strip():
            raise ValueError("No hay texto para analizar")
        
        analysis = {
            "timestamp": datetime.now().isoformat(),
            "text_length": len(text),
            "word_count": len(text.split())
        }
        
        if self.config.get("ai_analysis_enabled"):
            # Clasificación
            analysis["classification"] = await self._classify_text(text)
            
            # Extracción de entidades
            analysis["entities"] = await self._extract_entities_from_text(text)
            
            # Resumen
            analysis["summary"] = await self._generate_summary(text)
            
            # Análisis de sentimiento
            analysis["sentiment"] = await self._analyze_sentiment(text)
            
            # Detección de idioma
            analysis["language"] = await self._detect_language(text)
        
        return analysis
    
    async def _classify_document(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Clasificar tipo de documento"""
        text = data.get("text", "")
        file_name = data.get("file_name", "")
        
        # Clasificación basada en reglas simples
        basic_classification = self._classify_by_patterns(text, file_name)
        
        # Clasificación con IA si está habilitada
        if self.config.get("ai_analysis_enabled"):
            ai_classification = await self._classify_text(text)
            return {
                "basic_classification": basic_classification,
                "ai_classification": ai_classification,
                "confidence": ai_classification.get("confidence", 0)
            }
        
        return {"classification": basic_classification}
    
    def _classify_by_patterns(self, text: str, file_name: str) -> Dict[str, Any]:
        """Clasificación básica usando patrones de texto"""
        text_lower = text.lower()
        name_lower = file_name.lower()
        
        # Patrones para diferentes tipos de documentos
        patterns = {
            "factura": ["factura", "invoice", "total:", "iva", "subtotal", "importe"],
            "contrato": ["contrato", "contract", "cláusula", "firma", "partes contratantes"],
            "cv": ["curriculum", "experiencia laboral", "educación", "habilidades"],
            "informe": ["informe", "report", "análisis", "conclusiones", "recomendaciones"],
            "email": ["de:", "para:", "asunto:", "from:", "to:", "subject:"],
            "legal": ["artículo", "ley", "decreto", "resolución", "jurisprudencia"]
        }
        
        scores = {}
        for doc_type, keywords in patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower or keyword in name_lower)
            scores[doc_type] = score
        
        # Encontrar el tipo con mayor puntuación
        if scores:
            best_type = max(scores, key=scores.get)
            confidence = scores[best_type] / len(patterns[best_type])
            
            return {
                "type": best_type,
                "confidence": min(confidence, 1.0),
                "scores": scores
            }
        
        return {"type": "desconocido", "confidence": 0, "scores": scores}
    
    # 🔍 Métodos auxiliares con IA
    
    async def _classify_text(self, text: str) -> Dict[str, Any]:
        """Clasificar texto usando IA"""
        prompt = f"""
        Analiza el siguiente texto y clasifícalo en una de estas categorías:
        - factura: facturas, recibos, documentos de compra
        - contrato: contratos, acuerdos, términos legales  
        - cv: curriculum vitae, hojas de vida
        - informe: informes, reportes, análisis
        - email: correos electrónicos, comunicaciones
        - legal: documentos legales, normativas
        - otros: cualquier otro tipo
        
        Texto: {text[:1000]}...
        
        Responde solo con: tipo|confianza (0-1)
        """
        
        response = await self.ai_service.generate_text(prompt)
        
        try:
            parts = response.strip().split("|")
            return {
                "type": parts[0],
                "confidence": float(parts[1]) if len(parts) > 1 else 0.5
            }
        except:
            return {"type": "otros", "confidence": 0.1}
    
    async def _extract_entities_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extraer entidades del texto usando IA"""
        prompt = f"""
        Extrae las siguientes entidades del texto:
        - PERSONA: nombres de personas
        - EMPRESA: nombres de empresas u organizaciones  
        - FECHA: fechas mencionadas
        - DINERO: cantidades monetarias
        - LUGAR: ubicaciones, direcciones
        
        Texto: {text[:2000]}...
        
        Formato: TIPO: entidad
        """
        
        response = await self.ai_service.generate_text(prompt)
        
        entities = []
        for line in response.split('\n'):
            if ':' in line:
                try:
                    entity_type, entity_value = line.split(':', 1)
                    entities.append({
                        "type": entity_type.strip(),
                        "value": entity_value.strip()
                    })
                except:
                    continue
        
        return entities
    
    async def _generate_summary(self, text: str) -> str:
        """Generar resumen del texto"""
        if len(text) < 200:
            return text  # Texto muy corto, no necesita resumen
        
        prompt = f"""
        Resume el siguiente texto en máximo 3 frases, capturando los puntos más importantes:
        
        {text[:3000]}...
        """
        
        return await self.ai_service.generate_text(prompt)
    
    async def _analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Analizar sentimiento del texto"""
        prompt = f"""
        Analiza el sentimiento del siguiente texto.
        Responde solo con: positivo|neutro|negativo|confianza (0-1)
        
        Texto: {text[:1000]}...
        """
        
        response = await self.ai_service.generate_text(prompt)
        
        try:
            parts = response.strip().split("|")
            return {
                "sentiment": parts[0],
                "confidence": float(parts[1]) if len(parts) > 1 else 0.5
            }
        except:
            return {"sentiment": "neutro", "confidence": 0.1}
    
    async def _detect_language(self, text: str) -> str:
        """Detectar idioma del texto"""
        # Detección simple por patrones comunes
        spanish_patterns = ["el ", "la ", "de ", "que ", "en ", "un ", "es ", "se ", "no ", "te "]
        english_patterns = ["the ", "and ", "to ", "of ", "a ", "in ", "is ", "it ", "you ", "that "]
        
        text_lower = text.lower()
        
        spanish_score = sum(1 for pattern in spanish_patterns if pattern in text_lower)
        english_score = sum(1 for pattern in english_patterns if pattern in text_lower)
        
        if spanish_score > english_score:
            return "español"
        elif english_score > spanish_score:
            return "inglés"
        else:
            return "desconocido"
    
    # 📊 Métodos de procesamiento por lotes
    
    async def _process_batch(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Procesar múltiples documentos en lote"""
        file_paths = data.get("file_paths", [])
        
        results = {
            "processed": 0,
            "failed": 0,
            "total": len(file_paths),
            "documents": [],
            "summary": {}
        }
        
        for file_path in file_paths:
            try:
                # Procesar cada documento
                doc_result = await self._extract_text({"file_path": file_path})
                
                # Análisis básico
                if self.config.get("auto_classify"):
                    classification = await self._classify_document({
                        "text": doc_result["text"],
                        "file_name": doc_result["file_name"]
                    })
                    doc_result["classification"] = classification
                
                results["documents"].append(doc_result)
                results["processed"] += 1
                
            except Exception as e:
                self.logger.error(f"❌ Error procesando {file_path}: {str(e)}")
                results["failed"] += 1
                results["documents"].append({
                    "file_path": file_path,
                    "error": str(e)
                })
        
        # Generar resumen del lote
        results["summary"] = self._generate_batch_summary(results["documents"])
        
        return results
    
    def _generate_batch_summary(self, documents: List[Dict]) -> Dict[str, Any]:
        """Generar resumen de procesamiento por lotes"""
        successful_docs = [doc for doc in documents if "error" not in doc]
        
        if not successful_docs:
            return {"total_words": 0, "document_types": {}, "languages": {}}
        
        total_words = sum(doc.get("word_count", 0) for doc in successful_docs)
        
        # Contar tipos de documentos
        doc_types = {}
        languages = {}
        
        for doc in successful_docs:
            # Tipos de documentos
            if "classification" in doc:
                doc_type = doc["classification"].get("classification", {}).get("type", "desconocido")
                doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
            
            # Idiomas
            lang = doc.get("language", "desconocido")
            languages[lang] = languages.get(lang, 0) + 1
        
        return {
            "total_words": total_words,
            "average_words": total_words // len(successful_docs) if successful_docs else 0,
            "document_types": doc_types,
            "languages": languages,
            "successful_documents": len(successful_docs),
            "total_pages": sum(doc.get("pages", 0) for doc in successful_docs)
        }
    
    # 📈 Métodos de estadísticas específicas
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Obtener estadísticas específicas del procesamiento"""
        base_metrics = self.get_status()["metrics"]
        
        return {
            **base_metrics,
            "documents_processed": self.processed_documents,
            "pages_extracted": self.extracted_pages,
            "ocr_operations": self.ocr_operations,
            "avg_pages_per_doc": round(self.extracted_pages / max(1, self.processed_documents), 2)
        }

# 🏭 Registrar el agente en el factory
from .base_agent import AgentFactory
AgentFactory.register("document_processor", DocumentProcessorAgent)